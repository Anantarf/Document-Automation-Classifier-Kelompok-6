
# app/services/parser_docx.py
"""
Ekstraksi teks dari DOCX menggunakan python-docx:
- Mengambil teks paragraf
- Mengambil teks dalam tabel (opsional)
"""

from docx import Document

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    texts = []

    # Paragraf
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)

    # Tabel (opsional)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)

    return "\n".join(texts)
